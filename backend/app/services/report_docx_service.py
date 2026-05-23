"""Sprint 9 Phase 4 — DOCX report exporter.

Mirrors :mod:`report_export_service` (PDF) but emits a Microsoft Word
document via ``python-docx``. Same content sections, same data, same
ordering — so users can pick whichever format their downstream
toolchain prefers without information loss.

Optional Phase 4.2: when an AI narrative is supplied via
``include_narrative=True`` and an LLM is configured, an executive-summary
paragraph is prepended just below the cover.
"""

from __future__ import annotations

import io
import logging
from collections import defaultdict
from datetime import datetime, timezone

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from sqlalchemy.orm import Session

from app.models.project import Project
from app.services.project_calc_service import run_project_calculation

logger = logging.getLogger(__name__)


def _set_table_header(row, font_size: int = 10) -> None:
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(font_size)


def _add_kv_table(doc, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate(rows):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
        for cell in table.rows[i].cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def export_valuation_docx(
    project_id: int,
    db: Session,
    *,
    narrative: str | None = None,
) -> bytes:
    """Render a Word .docx valuation report and return it as bytes."""
    project = db.query(Project).filter(Project.id == project_id).first()
    if project is None:
        raise ValueError(f"Project {project_id} not found")

    summary, line_results = run_project_calculation(project_id=project_id, db=db)

    doc = Document()

    # ── Page setup: A4 with reasonable margins ──
    for section in doc.sections:
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)

    # ── Title ──
    title = doc.add_heading("工程计价报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Project info ──
    info_rows: list[tuple[str, str]] = [
        ("项目名称", project.name),
        ("所在地区", project.region),
        ("项目类型", project.project_type or ""),
        ("计价标准", project.standard_type or ""),
        ("币种", project.currency or ""),
    ]
    if project.description:
        info_rows.append(("项目描述", (project.description or "")[:200]))
    if project.budget:
        info_rows.append(("预算", f"{project.budget:,.2f}"))
    info_rows.append(
        ("报告时间", datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC"))
    )
    _add_kv_table(doc, info_rows)
    doc.add_paragraph()

    # ── AI narrative (optional Phase 4.2) ──
    if narrative:
        doc.add_heading("摘要", level=1)
        for para in narrative.strip().split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
        doc.add_paragraph()

    # ── Cost summary ──
    doc.add_heading("一、费用汇总", level=1)
    cost_rows = [
        ("直接费", f"{summary.total_direct:,.2f}"),
        ("管理费", f"{summary.total_management:,.2f}"),
        ("利润", f"{summary.total_profit:,.2f}"),
        ("规费", f"{summary.total_regulatory:,.2f}"),
        ("税前合计", f"{summary.total_pre_tax:,.2f}"),
        ("税金", f"{summary.total_tax:,.2f}"),
        ("措施费", f"{summary.total_measures:,.2f}"),
        ("工程总价", f"{summary.grand_total:,.2f}"),
    ]
    cost_table = doc.add_table(rows=1 + len(cost_rows), cols=2)
    cost_table.style = "Light Grid Accent 1"
    cost_table.rows[0].cells[0].text = "费用项目"
    cost_table.rows[0].cells[1].text = "金额"
    _set_table_header(cost_table.rows[0])
    for i, (k, v) in enumerate(cost_rows, start=1):
        cost_table.rows[i].cells[0].text = k
        cost_table.rows[i].cells[1].text = v
    # Highlight grand total row
    last_cells = cost_table.rows[-1].cells
    for cell in last_cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    doc.add_paragraph()

    # ── Division breakdown ──
    doc.add_heading("二、分部工程汇总", level=1)
    div_totals: dict[str, float] = defaultdict(float)
    div_counts: dict[str, int] = defaultdict(int)
    for boq, result in line_results:
        div = boq.division or "未分类"
        div_totals[div] += result.total
        div_counts[div] += 1

    grand = summary.grand_total or 1
    div_table = doc.add_table(rows=1 + len(div_totals) + 1, cols=4)
    div_table.style = "Light Grid Accent 1"
    headers = ("分部工程", "清单数", "合计金额", "占比")
    for j, h in enumerate(headers):
        div_table.rows[0].cells[j].text = h
    _set_table_header(div_table.rows[0])
    for i, (div, total) in enumerate(
        sorted(div_totals.items(), key=lambda x: -x[1]), start=1
    ):
        cells = div_table.rows[i].cells
        cells[0].text = div
        cells[1].text = str(div_counts[div])
        cells[2].text = f"{total:,.2f}"
        cells[3].text = f"{total / grand * 100:.1f}%"
    # total row
    last_idx = len(div_totals) + 1
    cells = div_table.rows[last_idx].cells
    cells[0].text = "合计"
    cells[1].text = str(sum(div_counts.values()))
    cells[2].text = f"{sum(div_totals.values()):,.2f}"
    cells[3].text = "100.0%"
    for cell in cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    doc.add_paragraph()

    # ── Line items ──
    doc.add_heading("三、分部分项工程计价表", level=1)
    line_table = doc.add_table(rows=1 + len(line_results) + 1, cols=7)
    line_table.style = "Light Grid Accent 1"
    headers = ("序号", "编码", "名称", "单位", "工程量", "综合单价", "合价")
    for j, h in enumerate(headers):
        line_table.rows[0].cells[j].text = h
    _set_table_header(line_table.rows[0], font_size=9)
    for idx, (boq, result) in enumerate(line_results, start=1):
        unit_price = result.total / boq.quantity if boq.quantity else 0
        cells = line_table.rows[idx].cells
        cells[0].text = str(idx)
        cells[1].text = boq.code or ""
        cells[2].text = (boq.name or "")[:30]
        cells[3].text = boq.unit or ""
        cells[4].text = f"{boq.quantity:,.2f}"
        cells[5].text = f"{unit_price:,.2f}"
        cells[6].text = f"{result.total:,.2f}"

    # Total row
    total_idx = len(line_results) + 1
    cells = line_table.rows[total_idx].cells
    cells[2].text = "合计"
    cells[6].text = f"{sum(r.total for _, r in line_results):,.2f}"
    for cell in cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    # ── Footer ──
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run(
        f"{project.name} — 工程计价报告 · 由 AI Native Valuation 系统生成"
    )
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
